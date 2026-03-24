"""
Document parser — extracts raw text from PDF, DOCX, and TXT files.
"""

import logging
from pathlib import Path
from typing import Dict, Any

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =========================
# OPTIONAL DEPENDENCIES
# =========================

try:
    from PyPDF2 import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("Install PyPDF2")

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("Install pytesseract pdf2image")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Install python-docx")

try:
    from tika import parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.warning("Install tika")


# =========================
# EXTRACTORS
# =========================

def extract_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF. Returns (text, page_count)."""

    if not PYPDF_AVAILABLE:
        return "", 0

    try:
        reader = PdfReader(file_path)
        text = ""

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"
            elif OCR_AVAILABLE:
                images = convert_from_path(file_path, first_page=i + 1, last_page=i + 1)
                if images:
                    text += pytesseract.image_to_string(images[0])

        return text, len(reader.pages)

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return "", 0


def extract_docx(file_path: str) -> tuple[str, int]:
    """Extract text from a DOCX. Returns (text, page_count=0)."""

    if not DOCX_AVAILABLE:
        return "", 0

    try:
        doc = Document(file_path)
        text = ""

        for p in doc.paragraphs:
            text += p.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"

        return text, 0

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return "", 0


def extract_txt(file_path: str) -> tuple[str, int]:
    """Extract text from a TXT file. Returns (text, page_count=0)."""

    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), 0
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        return "", 0


def extract_with_tika(file_path: str) -> str:
    """Fallback extraction using Apache Tika."""

    if not TIKA_AVAILABLE:
        return ""

    try:
        raw = parser.from_file(file_path)
        if raw and raw.get("content"):
            return raw["content"]
    except Exception as e:
        logger.error(f"Tika failed: {e}")

    return ""


# =========================
# UNIVERSAL PARSER
# =========================

def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Parse any supported document and return raw text + metadata.

    Returns:
        {"success": bool, "text": str, "page_count": int, "error": str}
    """

    path = Path(file_path)

    if not path.exists():
        return {"success": False, "text": "", "page_count": 0, "error": "file not found"}

    text = ""
    page_count = 0

    ext = path.suffix.lower()

    if ext == ".pdf":
        text, page_count = extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text, page_count = extract_docx(file_path)
    elif ext == ".txt":
        text, page_count = extract_txt(file_path)

    # Fallback to Tika
    if not text:
        text = extract_with_tika(file_path)

    if not text:
        return {"success": False, "text": "", "page_count": 0, "error": "text extraction failed"}

    return {"success": True, "text": text, "page_count": page_count}
