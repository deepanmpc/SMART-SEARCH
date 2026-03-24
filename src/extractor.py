import logging
import re
from pathlib import Path
from typing import Dict, Any, List

# =========================
# Logging
# =========================

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =========================
# Dependencies
# =========================

try:
    from PyPDF2 import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("Install PyPDF2")

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("Install pytesseract pdf2image")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Install python-docx")

try:
    from tika import parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.warning("Install tika")


# =========================
# TEXT CHUNKING
# =========================

def normalize_text(text: str) -> str:
    """
    Clean raw extracted text before chunking.

    Fixes:
    - Hyphenated line-breaks  e.g. "detec-\ntion" → "detection"
    - Mid-word newlines       e.g. "feature\nextraction" → "feature extraction"
    - Tabs / non-breaking spaces → regular space
    - Multiple blank lines    → single blank line (preserve paragraph signal)
    - Leading/trailing whitespace
    """

    # 1. Rejoin words that were split across lines with a hyphen
    text = re.sub(r"-\n(\S)", r"\1", text)

    # 2. Replace tabs and non-breaking spaces with a regular space
    text = re.sub(r"[ \t\xa0]+", " ", text)

    # 3. Join lines that are NOT paragraph breaks
    #    (A real paragraph break = blank line; a soft wrap = single \n)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)

    # 4. Collapse runs of 3+ newlines down to 2 (one blank line)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 5. Strip leftover page-marker noise (e.g. "--- Page 1 ---")
    text = re.sub(r"-{2,}\s*Page\s*\d+\s*-{2,}", "", text, flags=re.IGNORECASE)

    return text.strip()


def chunk_text(text: str, chunk_size: int = 120, overlap: int = 30) -> List[str]:
    """
    Recursive character chunker for semantic search pipelines.

    Pre-processes text with normalize_text() first, then splits using
    the following strategy (coarsest → finest):
      1. Paragraph breaks (double newlines)
      2. Sentence endings (. ? !)
      4. Sliding-window over words (final fallback)

    Chunks smaller than 30 characters are silently dropped.
    """

    text = normalize_text(text)

    # ---------- helpers ----------

    def _word_count(s: str) -> int:
        return len(s.split())

    def _sliding_window(text: str) -> List[str]:
        """Fixed-size sliding window — the final fallback."""
        words = text.split()
        result = []
        start = 0
        while start < len(words):
            chunk = " ".join(words[start : start + chunk_size])
            if len(chunk.strip()) > 30:
                result.append(chunk)
            start += chunk_size - overlap
        return result

    def _split_and_merge(text: str, separators: List[str]) -> List[str]:
        """
        Try each separator in order. Merge small pieces into chunks
        that stay at or below chunk_size words, then recurse on anything
        that is still too large.
        """
        if not separators:
            return _sliding_window(text)

        sep, *rest = separators
        parts = [p.strip() for p in re.split(sep, text) if p.strip()]

        # If the separator produced no useful split, try the next one
        if len(parts) <= 1:
            return _split_and_merge(text, rest)

        chunks: List[str] = []
        current: List[str] = []
        current_words = 0

        for part in parts:
            part_words = _word_count(part)

            # Part alone is bigger than target — recurse deeper
            if part_words > chunk_size:
                if current:
                    merged = " ".join(current)
                    if len(merged.strip()) > 30:
                        chunks.append(merged)
                    current = []
                    current_words = 0
                chunks.extend(_split_and_merge(part, rest))
                continue

            # Adding this part would exceed target — flush and carry overlap
            if current_words + part_words > chunk_size and current:
                merged = " ".join(current)
                if len(merged.strip()) > 30:
                    chunks.append(merged)
                overlap_words = " ".join(current).split()[-overlap:]
                current = overlap_words + [part]
                current_words = len(current)
            else:
                current.append(part)
                current_words += part_words

        # Flush remainder
        if current:
            merged = " ".join(current)
            if len(merged.strip()) > 30:
                chunks.append(merged)

        return chunks

    # ---------- main ----------

    separators = [
        r"\n\n+",          # paragraph breaks
        r"(?<=[.?!])\s+",  # sentence endings
    ]

    return _split_and_merge(text, separators)


# =========================
# PDF EXTRACTION
# =========================

def extract_pdf(file_path: str) -> str:

    if not PYPDF_AVAILABLE:
        return ""

    try:

        reader = PdfReader(file_path)
        text = ""

        for i, page in enumerate(reader.pages):

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

            elif OCR_AVAILABLE:

                images = convert_from_path(
                    file_path,
                    first_page=i + 1,
                    last_page=i + 1
                )

                if images:
                    text += pytesseract.image_to_string(images[0])

        return text

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


# =========================
# DOCX EXTRACTION
# =========================

def extract_docx(file_path: str) -> str:

    if not DOCX_AVAILABLE:
        return ""

    try:

        doc = Document(file_path)
        text = ""

        for p in doc.paragraphs:
            text += p.text + "\n"

        for table in doc.tables:

            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"

        return text

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


# =========================
# TXT EXTRACTION
# =========================

def extract_txt(file_path: str) -> str:

    try:

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        return ""


# =========================
# TIKA FALLBACK
# =========================

def extract_with_tika(file_path: str) -> str:

    if not TIKA_AVAILABLE:
        return ""

    try:

        raw = parser.from_file(file_path)

        if raw and raw.get("content"):
            return raw["content"]

    except Exception as e:
        logger.error(f"Tika failed: {e}")

    return ""


# =========================
# UNIVERSAL EXTRACTOR
# =========================

def extract_any_document(file_path: str) -> Dict[str, Any]:

    path = Path(file_path)

    if not path.exists():

        return {
            "success": False,
            "error": "file not found"
        }

    text = ""

    # Select extractor
    if path.suffix.lower() == ".pdf":
        text = extract_pdf(file_path)

    elif path.suffix.lower() in [".docx", ".doc"]:
        text = extract_docx(file_path)

    elif path.suffix.lower() == ".txt":
        text = extract_txt(file_path)

    # Fallback to Tika
    if not text:
        text = extract_with_tika(file_path)

    if not text:

        return {
            "success": False,
            "error": "text extraction failed"
        }

    # Chunk the text
    chunks = chunk_text(text)

    return {
        "success": True,
        "file_path": file_path,
        "chunks": chunks
    }


# =========================
# CLI TEST
# =========================

if __name__ == "__main__":

    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <file>")
        exit()

    result = extract_any_document(sys.argv[1])

    if result["success"]:

        print(f"Extracted chunks: {len(result['chunks'])}")

        for c in result["chunks"][:2]:
            print("\nChunk preview:\n", c[:200])

    else:

        print("Error:", result["error"])